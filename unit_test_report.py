#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
单元测试文档生成器
用于生成包含单元测试结果的Word文档
"""

import os
import datetime
import unittest
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL


class TestResult:
    """测试结果数据类"""
    def __init__(self, name, status, execution_time, details=None):
        self.name = name
        self.status = status  # "通过", "失败", "错误", "跳过"
        self.execution_time = execution_time  # 单位：秒
        self.details = details or ""


class UnitTestDocumentGenerator:
    """单元测试文档生成器"""
    
    def __init__(self, project_name, version, author):
        """初始化文档生成器
        
        Args:
            project_name: 项目名称
            version: 版本号
            author: 作者
        """
        self.project_name = project_name
        self.version = version
        self.author = author
        self.test_results = []
        self.document = Document()
        self.date = datetime.datetime.now().strftime("%Y-%m-%d")
        
    def add_test_result(self, test_result):
        """添加测试结果
        
        Args:
            test_result: TestResult对象
        """
        self.test_results.append(test_result)
    
    def add_test_results_from_unittest(self, test_suite):
        """从unittest测试套件添加测试结果
        
        Args:
            test_suite: unittest.TestSuite对象
        """
        class SuccessTrackingResult(unittest.TestResult):
            def __init__(self):
                super().__init__()
                self.successes = []
            
            def addSuccess(self, test):
                super().addSuccess(test)
                self.successes.append(test)
        
        result = SuccessTrackingResult()
        start_time = datetime.datetime.now()
        test_suite.run(result)
        end_time = datetime.datetime.now()
        total_execution_time = (end_time - start_time).total_seconds()
        avg_time_per_test = total_execution_time / (len(result.successes) + len(result.failures) + len(result.errors) + len(result.skipped)) if (len(result.successes) + len(result.failures) + len(result.errors) + len(result.skipped)) > 0 else 0
        
        # 处理通过的测试
        for test in result.successes:
            test_name = str(test)
            execution_time = avg_time_per_test  # 使用平均执行时间作为近似值
            self.add_test_result(TestResult(test_name, "通过", execution_time))
        
        # 处理失败的测试
        for test, error_msg in result.failures:
            test_name = str(test)
            execution_time = avg_time_per_test  # 使用平均执行时间作为近似值
            self.add_test_result(TestResult(test_name, "失败", execution_time, error_msg))
        
        # 处理错误的测试
        for test, error_msg in result.errors:
            test_name = str(test)
            execution_time = avg_time_per_test  # 使用平均执行时间作为近似值
            self.add_test_result(TestResult(test_name, "错误", execution_time, error_msg))
        
        # 处理跳过的测试
        for test, reason in result.skipped:
            test_name = str(test)
            execution_time = 0
            self.add_test_result(TestResult(test_name, "跳过", execution_time, reason))
    
    def _add_title(self):
        """添加文档标题"""
        title = self.document.add_heading(f'{self.project_name} 单元测试报告', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加副标题
        subtitle = self.document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run(f'版本: {self.version} | 日期: {self.date} | 作者: {self.author}')
        subtitle_run.font.size = Pt(12)
        subtitle_run.font.italic = True
        
        self.document.add_paragraph()
    
    def _add_summary(self):
        """添加测试摘要"""
        self.document.add_heading('测试摘要', level=1)
        
        # 计算统计数据
        total_tests = len(self.test_results)
        passed_tests = sum(1 for result in self.test_results if result.status == "通过")
        failed_tests = sum(1 for result in self.test_results if result.status == "失败")
        error_tests = sum(1 for result in self.test_results if result.status == "错误")
        skipped_tests = sum(1 for result in self.test_results if result.status == "跳过")
        
        if total_tests > 0:
            pass_rate = (passed_tests / total_tests) * 100
        else:
            pass_rate = 0
        
        # 创建摘要表格
        table = self.document.add_table(rows=2, cols=6)
        table.style = 'Table Grid'
        
        # 表头
        header_cells = table.rows[0].cells
        header_cells[0].text = '总测试数'
        header_cells[1].text = '通过'
        header_cells[2].text = '失败'
        header_cells[3].text = '错误'
        header_cells[4].text = '跳过'
        header_cells[5].text = '通过率'
        
        # 设置表头格式
        for cell in header_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.bold = True
        
        # 数据行
        data_cells = table.rows[1].cells
        data_cells[0].text = str(total_tests)
        data_cells[1].text = str(passed_tests)
        data_cells[2].text = str(failed_tests)
        data_cells[3].text = str(error_tests)
        data_cells[4].text = str(skipped_tests)
        data_cells[5].text = f'{pass_rate:.2f}%'
        
        # 设置数据行格式
        for cell in data_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 根据通过率设置颜色
        if pass_rate >= 90:
            data_cells[5].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)  # 绿色
        elif pass_rate >= 70:
            data_cells[5].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 165, 0)  # 橙色
        else:
            data_cells[5].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)  # 红色
        
        self.document.add_paragraph()
    
    def _add_test_details(self):
        """添加测试详情"""
        self.document.add_heading('测试详情', level=1)
        
        if not self.test_results:
            self.document.add_paragraph('没有测试结果可显示。')
            return
        
        # 创建测试详情表格
        table = self.document.add_table(rows=len(self.test_results) + 1, cols=4)
        table.style = 'Table Grid'
        
        # 表头
        header_cells = table.rows[0].cells
        header_cells[0].text = '测试名称'
        header_cells[1].text = '状态'
        header_cells[2].text = '执行时间 (秒)'
        header_cells[3].text = '详细信息'
        
        # 设置表头格式
        for cell in header_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.bold = True
        
        # 数据行
        for i, result in enumerate(self.test_results, 1):
            row_cells = table.rows[i].cells
            row_cells[0].text = result.name
            row_cells[1].text = result.status
            row_cells[2].text = f'{result.execution_time:.3f}'
            row_cells[3].text = result.details
            
            # 设置状态单元格的颜色
            status_run = row_cells[1].paragraphs[0].runs[0]
            if result.status == "通过":
                status_run.font.color.rgb = RGBColor(0, 128, 0)  # 绿色
            elif result.status == "失败":
                status_run.font.color.rgb = RGBColor(255, 0, 0)  # 红色
            elif result.status == "错误":
                status_run.font.color.rgb = RGBColor(255, 0, 0)  # 红色
            elif result.status == "跳过":
                status_run.font.color.rgb = RGBColor(128, 128, 128)  # 灰色
        
        self.document.add_paragraph()
    
    def _add_conclusion(self):
        """添加结论部分"""
        self.document.add_heading('结论', level=1)
        
        # 计算统计数据
        total_tests = len(self.test_results)
        passed_tests = sum(1 for result in self.test_results if result.status == "通过")
        
        if total_tests > 0:
            pass_rate = (passed_tests / total_tests) * 100
        else:
            pass_rate = 0
        
        paragraph = self.document.add_paragraph()
        if pass_rate == 100:
            paragraph.add_run('所有测试均已通过，代码质量良好。').bold = True
        elif pass_rate >= 90:
            paragraph.add_run('大部分测试已通过，但仍有少量问题需要修复。').bold = True
        elif pass_rate >= 70:
            paragraph.add_run('测试通过率一般，需要关注失败的测试用例并进行修复。').bold = True
        else:
            paragraph.add_run('测试通过率较低，代码可能存在严重问题，需要全面检查。').bold = True
    
    def generate(self, output_path):
        """生成单元测试文档
        
        Args:
            output_path: 输出文件路径
        
        Returns:
            bool: 是否成功生成文档
        """
        try:
            # 添加文档各部分
            self._add_title()
            self._add_summary()
            self._add_test_details()
            self._add_conclusion()
            
            # 保存文档
            self.document.save(output_path)
            return True
        except Exception as e:
            print(f"生成文档时出错: {e}")
            return False


def demo():
    """演示如何使用单元测试文档生成器"""
    # 创建文档生成器
    generator = UnitTestDocumentGenerator(
        project_name="示例项目",
        version="1.0.0",
        author="测试团队"
    )
    
    # 添加一些示例测试结果
    generator.add_test_result(TestResult(
        name="test_login_success",
        status="通过",
        execution_time=0.123
    ))
    
    generator.add_test_result(TestResult(
        name="test_login_invalid_password",
        status="通过",
        execution_time=0.056
    ))
    
    generator.add_test_result(TestResult(
        name="test_user_registration",
        status="失败",
        execution_time=0.089,
        details="AssertionError: 预期用户数量增加1，但实际没有变化"
    ))
    
    generator.add_test_result(TestResult(
        name="test_password_reset",
        status="错误",
        execution_time=0.045,
        details="TypeError: send_email() missing 1 required positional argument: 'recipient'"
    ))
    
    generator.add_test_result(TestResult(
        name="test_admin_access",
        status="跳过",
        execution_time=0,
        details="需要管理员权限配置"
    ))
    
    # 生成文档
    output_path = "单元测试报告.docx"
    success = generator.generate(output_path)
    
    if success:
        print(f"单元测试文档已生成: {os.path.abspath(output_path)}")
    else:
        print("生成单元测试文档失败")


if __name__ == "__main__":
    demo() 